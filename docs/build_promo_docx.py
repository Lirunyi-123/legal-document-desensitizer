#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 简介.docx（launch_messaging_guide 预设）。"""

import os

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- 预设 token（launch_messaging_guide = compact_reference_guide） ----------
EA_FONT = 'Heiti SC'          # eastAsia 字体（本机已装，渲染保真）
ASCII_FONT = 'Calibri'        # 拉丁字体
CODE_FONT = 'Consolas'
INK_BLUE = RGBColor(0x0B, 0x25, 0x45)
H1_BLUE = RGBColor(0x2E, 0x74, 0xB5)
H3_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x55, 0x55, 0x55)
TABLE_HEADER_FILL = 'E8EEF5'
CODE_FILL = 'F2F4F7'
CALLOUT_FILL = 'F4F6F9'
USABLE_DXA = 9360
TBL_IND_DXA = 120


def set_run_font(run, size=11, bold=False, color=None, font=ASCII_FONT,
                 ea=EA_FONT, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:eastAsia'), ea)


def para(doc, text='', size=11, bold=False, color=None, before=0, after=6,
         line=300, align=None, style=None, italic=False):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line / 240.0
    if align is not None:
        pf.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def shade_paragraph(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def border_paragraph(p, color='D0D7E2'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '6')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)


def rich_para(doc, parts, before=0, after=6, line=300, indent=None):
    """parts: [(text, dict(size/bold/color/italic)), ...]"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line / 240.0
    if indent is not None:
        pf.left_indent = Inches(indent)
    for text, fmt in parts:
        r = p.add_run(text)
        set_run_font(r, **fmt)
    return p


# ---------- 列表：真实 numbering ----------
def add_numbering(doc, bullet=True, start=1):
    """创建一个 bullet 或 decimal 的抽象编号定义，返回 numId。"""
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    # 找最大 abstractNumId / numId
    abs_ids = [int(a.get(qn('w:abstractNumId'))) for a in numbering.findall(qn('w:abstractNum'))]
    num_ids = [int(n.get(qn('w:numId'))) for n in numbering.findall(qn('w:num'))]
    abs_id = (max(abs_ids) + 1) if abs_ids else 0
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement('w:abstractNum')
    abstract.set(qn('w:abstractNumId'), str(abs_id))
    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), '0')
    start_el = OxmlElement('w:start')
    start_el.set(qn('w:val'), str(start))
    lvl.append(start_el)
    numfmt = OxmlElement('w:numFmt')
    numfmt.set(qn('w:val'), 'bullet' if bullet else 'decimal')
    lvl.append(numfmt)
    lvlText = OxmlElement('w:lvlText')
    lvlText.set(qn('w:val'), '\u2022' if bullet else '%1.')
    lvl.append(lvlText)
    lvlJc = OxmlElement('w:lvlJc')
    lvlJc.set(qn('w:val'), 'left')
    lvl.append(lvlJc)
    pPr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '540')       # 0.375in
    ind.set(qn('w:hanging'), '270')    # 0.187in marker
    pPr.append(ind)
    lvl.append(pPr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    absRef = OxmlElement('w:abstractNumId')
    absRef.set(qn('w:val'), str(abs_id))
    num.append(absRef)
    numbering.append(num)
    return num_id


def list_item(doc, num_id, text='', bold_lead=None, size=11, after=4):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numIdEl = OxmlElement('w:numId')
    numIdEl.set(qn('w:val'), str(num_id))
    numPr.append(ilvl)
    numPr.append(numIdEl)
    pPr.append(numPr)
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.line_spacing = 300 / 240.0
    if bold_lead:
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=size, bold=True)
    if text:
        r2 = p.add_run(text)
        set_run_font(r2, size=size)
    return p


def code_block(doc, lines):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(6 if i == 0 else 0)
        pf.space_after = Pt(6 if i == len(lines) - 1 else 0)
        pf.line_spacing = 1.0
        pf.left_indent = Inches(0.2)
        r = p.add_run(line)
        set_run_font(r, size=9.5, font=CODE_FONT)
        shade_paragraph(p, CODE_FILL)
    return p


# ---------- 表格 ----------
def make_table(doc, headers, rows, widths_in, header_fill=TABLE_HEADER_FILL):
    """固定几何表格：tblW/tblInd/tblGrid/tcW 全部显式。"""
    total_in = sum(widths_in)
    assert abs(total_in - 6.5) < 0.01, f'列宽合计须 6.5in，实际 {total_in}'
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(USABLE_DXA))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(TBL_IND_DXA))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    # 单元格边距
    tblCellMar = OxmlElement('w:tblCellMar')
    for side, val in (('top', 80), ('start', 120), ('bottom', 80), ('end', 120)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tblCellMar.append(el)
    tblPr.append(tblCellMar)

    grid = tbl.find(qn('w:tblGrid'))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement('w:tblGrid')
    for w in widths_in:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w * 1440)))
        grid.append(gc)
    tbl.insert(list(tbl).index(tblPr) + 1, grid)

    def fill_cell(cell, text, col, bold=False, color=None, size=10, align='left'):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tcPr = cell._tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(widths_in[col] * 1440)))
        tcPr.append(tcW)
        p = cell.paragraphs[0]
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.15
        pf.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER}[align]
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)

    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        fill_cell(cell, h, j, bold=True, size=10)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), header_fill)
        cell._tc.get_or_add_tcPr().append(shd)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            bold = False
            color = None
            if isinstance(val, tuple):
                val, bold, color = val
            fill_cell(table.cell(i, j), val, j, bold=bold, color=color)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_run_font(r, size={1: 16, 2: 13, 3: 12}[level],
                     bold=True, color=H1_BLUE if level < 3 else H3_BLUE)
    pf = p.paragraph_format
    pf.space_before = Pt({1: 18, 2: 14, 3: 10}[level])
    pf.space_after = Pt({1: 10, 2: 7, 3: 5}[level])
    return p


def add_header_footer(doc, label):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ''
    pf = hp.paragraph_format
    pf.space_after = Pt(2)
    r_label = hp.add_run(label + '　　')
    set_run_font(r_label, size=8.5, color=GRAY)
    hp.add_run().add_tab()
    # 右侧页码
    r = hp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar1)
    r._r.append(instr)
    r._r.append(fldChar2)
    for run in hp.runs:
        set_run_font(run, size=8.5, color=GRAY)
    # 右对齐 tab stop
    tabs = hp.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    # 页眉下细线
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), 'C9D3E0')
    pBdr.append(bottom)
    pPr.append(pBdr)


def main():
    doc = Document()
    # 页面
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for attr in ('top_margin', 'right_margin', 'bottom_margin', 'left_margin'):
        setattr(sec, attr, Inches(1))
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    # Normal 样式
    normal = doc.styles['Normal']
    normal.font.name = ASCII_FONT
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), EA_FONT)
    npf = normal.paragraph_format
    npf.space_after = Pt(6)
    npf.line_spacing = 300 / 240.0

    add_header_footer(doc, '法律文书脱敏工具 · 简介')

    # ---------- 标题块 ----------
    t = para(doc, '法律文书脱敏工具', size=24, bold=True, color=INK_BLUE,
             before=0, after=4)
    para(doc, 'Legal Document Desensitizer', size=12, color=GRAY, after=10)
    para(doc, '面向中国律师与法律实务的本地优先文书脱敏工具', size=13, bold=True,
         color=H1_BLUE, after=4)
    meta = '规则引擎 + 语义占位符 ｜ 两阶段工作流（审阅清单） ｜ 加密映射表 ｜ 一键无损还原 ｜ 红队评测量化'
    para(doc, meta, size=10, color=GRAY, after=12)

    # 核心承诺 callout
    cp = doc.add_paragraph()
    pf = cp.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(14)
    pf.line_spacing = 1.3
    r = cp.add_run('核心承诺：')
    set_run_font(r, size=11.5, bold=True, color=INK_BLUE)
    r2 = cp.add_run('结构化数据零依赖本地处理、可量化验证、可无损还原、不盲目信任模型输出。')
    set_run_font(r2, size=11.5)
    shade_paragraph(cp, CALLOUT_FILL)
    border_paragraph(cp)

    # ---------- 一句话介绍 ----------
    add_heading(doc, '一句话介绍', 1)
    para(doc, '把判决书、合同、聊天记录、证据材料里的敏感信息（身份证号、手机号、银行卡号、人名、'
              '地址、金额、案情细节……）自动替换为语义占位符（[当事人甲（原告）]、[身份证号]），'
              '同时生成加密映射表，需要原文时一键无损还原。')

    # ---------- 为什么需要它 ----------
    add_heading(doc, '为什么需要它', 1)
    para(doc, '律师日常工作离不开涉密文书：起诉状、判决书、合同、尽调报告、聊天记录。对外分享、'
              '上传 AI 辅助、交给协作方时，一条没脱敏的身份证号或当事人姓名就可能构成泄密。'
              '而通用工具不懂中国法律语境：要么漏掉案号、统一社会信用代码、律师执业证号这类'
              '中国特有实体，要么把“粤B88888”车牌当微信号、把日期当身份信息。')
    para(doc, '本工具专为中国法律文书设计，覆盖 20+ 类中国法律敏感实体，并针对律师真正关心的'
              '三件事做了设计：脱得干净、查得到底、还原得回来。')

    # ---------- 适用场景 ----------
    add_heading(doc, '适用场景', 1)
    bullet_id = add_numbering(doc, bullet=True)
    for s in ['判决书 / 裁定书 / 调解书的对外脱敏',
              '合同、协议中的当事人信息与商业条款脱敏',
              '微信聊天记录、邮件等沟通证据的提交前脱敏',
              '尽调材料、证据册的涉密信息筛查',
              '文书上传云端 AI 前的预处理（结构化号码本地先替换）',
              '律所内部归档 / 协作时的隐私合规',
              '格式预处理：docx / pdf / html / csv / 扫描件图片先经内置 format-converter '
              '转成 txt / md / json 再脱敏（判决书 PDF、网上复制的裁判文书、当事人名单等）']:
        list_item(doc, bullet_id, s)

    # ---------- 核心亮点 ----------
    add_heading(doc, '好用的地方（核心亮点）', 1)
    dec_id = add_numbering(doc, bullet=False)
    highlights = [
        ('中国法律实体全覆盖（规则层，无需模型）',
         '身份证号（GB 11643 校验码）、统一社会信用代码（GB 32100 校验码）、银行卡号（Luhn）、'
         '手机号、邮箱、微信号、律师执业证号、组织机构代码、护照/港澳通行证/驾驶证、案号、'
         '车牌号、出生日期、金额、角色人名、裸人名、公司名、地址（含无省市区层级的写法）。'),
        ('语义占位符，律师可读',
         '不打马赛克，替换为带角色语义的占位符：同一人物全文对应同一个占位符'
         '（实体归一化 + 种子传播），案情脉络完整保留。'),
        ('加密映射表 + 一键无损还原',
         'AES-256-GCM + PBKDF2 加密映射表，密钥不落终端；restore 一条命令还原，逐字节一致，'
         '多个 [金额] 也不会还原错位。'),
        ('红队评测，用数据说话',
         '内置 60 个红队用例（含真实判决书实战负样本），一键输出分类型召回率/误报/泄露报告；'
         '当前基线 60/60 通过，结构化期望 70 项召回率 100%，保留项误报 0，泄露 0。'),
        ('失败安全，不盲目信任模型',
         'LLM 层输出若增删行数、改动已有占位符、声称替换的值仍残留原文，整套输出被拒绝并中止。'),
        ('可复现、可自检',
         'selfcheck.py 一键验证文件、59 项单元测试、红队评测、还原往返，任何拷贝跑一遍即可比对一致性。'),
        ('两阶段工作流，无需接入任何外部 API',
         '阶段一（规则层）：关键信息本地必清，mask --review 一次生成脱敏文书 + 审阅清单，'
         '自动校验关键信息清零并列出剩余低优先级项；阶段二（语义层）：律师审阅后说'
         '"继续语义层脱敏"，配置了本技能的 AI 直接执行，无需 Ollama、无需 API Key；'
         '本地/云端模型仍为可选替代，不再是默认依赖。'),
        ('实战打磨，专治 OCR 扫描版文书',
         '以 42 页 OCR 扫描判决书做端到端验证：387/387 段无损还原（映射表逐字节一致，'
         '含数字与单位间带空格/尾随空格的写法）；角色词后名词不再误伤（提供担保/处签名/印章/私章）；'
         '公司名不吞上下文；OCR 空格兼容（公司简称内空格、金额数字与单位间空格）；金额支持无单位大数。'
         'OCR 识别最低限度（实测）：字号 ≥8pt 稳定识别；分辨率 ≥150 DPI（工具默认 200）；'
         '长度无下限（单个汉字也能识别）；真正的限制是清晰度而非大小，低清翻拍/传真识别率下降，'
         '审阅清单会兜底提示 not_found/残留。'),
        ('原图涂黑（v3.9）',
         '图片/截图输入直接在原图上涂黑敏感区域，保留原图版式输出 PDF：'
         'macOS Vision 带坐标 OCR 定位敏感值 → 原图对应区域涂黑 → 输出 PDF，'
         '双层 residual 校验（像素级确认涂黑 + OCR 复查）。实测真实建行流水截图'
         '76 处涂黑、覆盖页面 25.4%，关键敏感区像素全黑确认。'),
        ('内置格式转换 format-converter（v4.1）',
         '纯 JS 无外部引擎：docx（mammoth）/ pdf（pdfjs-dist）/ 文本互转（turndown、'
         'marked、csv-parse）/ 图片（sharp）全部本地处理，npm install 即可用。'
         '扩充脱敏输入链路：判决书 PDF → txt、网上复制的裁判文书 HTML → md、律师文书 '
         'docx → txt、当事人名单 csv → json、扫描件图片 → pdf。扫描件/图片型 PDF 无文本层'
         '时明确报错并指引走 OCR 通道，不静默产出空文件；刻意排除与脱敏无关的转换'
         '（音视频 / NCM / ZIP / Office 全格式），保持轻量。来源 FlyingMouse Format（MIT）。'),
    ]
    for lead, body in highlights:
        list_item(doc, dec_id, body, bold_lead=lead + '：')

    # 使用决策 callout（format-converter 单向提取，先转换还是直接脱敏）
    warn = doc.add_paragraph()
    wp = warn.paragraph_format
    wp.space_before = Pt(8)
    wp.space_after = Pt(12)
    wp.line_spacing = 1.25
    wr = warn.add_run('⚠️ 使用决策：format-converter 是单向提取器，转换后丢失版式且不支持转回 '
                      'docx/pdf（restore 只还原内容、不还原格式）。需要最终 docx/pdf 原格式 '
                      '（归档/庭审/提交）时不要先转换，直接 mask -f 合同.docx 或 '
                      'mask -f 判决书.pdf --pdf-redact（保格式、可还原）；只有需要提取文本 '
                      '（核对/摘录/喂 LLM）或输入是 html/csv/heic 等脱敏工具不直接支持的格式时，'
                      '才先用 format-converter 转成 txt/md/json。要"转回原格式"就别转，直接脱敏。')
    set_run_font(wr, size=10, bold=True, color=GRAY)
    shade_paragraph(warn, 'FFF4E6')
    border_paragraph(warn)

    # ---------- 如何运行 ----------
    add_heading(doc, '如何运行', 1)
    add_heading(doc, '安装', 2)
    code_block(doc, [
        'git clone https://github.com/Lirunyi-123/legal-document-desensitizer',
        'cd legal-document-desensitizer',
        'pip3 install -r requirements.txt   # jieba 必装；docx/pdf/加密为可选',
        'cd format-converter && npm install && cd ..   # 格式转换核心（可选，需 node >= 22.13）',
    ])
    add_heading(doc, '常用命令', 2)
    code_block(doc, [
        '# 格式预处理（format-converter，可选）：docx/pdf/html/csv/图片 → txt/md/json',
        'node format-converter/cli.js 判决书.pdf txt -o 判决书.txt',
        'node format-converter/cli.js 合同.docx txt -o 合同.txt',
        'node format-converter/cli.js 裁判文书.html md',
        'node format-converter/cli.js 当事人.csv json',
        '',
        'python3 desensitize.py scan -f 判决书.docx                  # 扫描敏感信息',
        'python3 desensitize.py mask -f 判决书.docx                  # 规则层脱敏（零模型）',
        'python3 desensitize.py mask -f 判决书.docx --review         # 阶段一：脱敏+审阅清单',
        'python3 desensitize.py mask -f 合同.docx --save-mapping 映射表.enc --encrypt-mapping',
        'python3 desensitize.py restore -f 合同_desensitized.docx -m 映射表.enc -o 还原.docx',
        'python3 desensitize.py full -f 判决书.docx --llm-model qwen2.5   # 可选：规则层+本地LLM',
        'python3 evaluate.py                                          # 红队评测',
        'python3 selfcheck.py                                         # 一键自检',
    ])
    add_heading(doc, '一键自检（推荐所有使用/宣传对象执行）', 2)
    para(doc, '运行 python3 selfcheck.py，输出全部 PASS 即证明拷贝与官方版本行为一致。')

    # ---------- 注意事项 ----------
    add_heading(doc, '注意事项（安全边界）', 1)
    add_heading(doc, '安全等级，请对号入座', 2)
    make_table(doc,
               ['操作', '替换了什么', '还剩什么', '能否上传云端 AI'],
               [
                   ['规则层 mask（阶段一）', '关键信息（身份证/手机/银行卡/案号等）+ 人名/公司名/地址/金额',
                    '法院名称、案情细节、公司简称、OCR 变体', '⚠️ 须先经审阅清单确认'],
                   ['规则层 + Agent 语义层（默认，无需外部 API）', '上述全部 + 语义层低优先级信息',
                    '无', '✅ 可以'],
                   ['规则层 + 本地 LLM（full）', '上述全部 + 案情细节/公司名', '无', '✅ 可以（模型在本地）'],
                   ['规则层 + 云端 API（full）', '上述全部', '无', '✅ 但 LLM 那一步数据会到服务商'],
               ],
               [1.65, 2.25, 1.35, 1.25])
    para(doc, '', size=4, after=2)
    add_heading(doc, '必须知道的三件事', 2)
    warn_id = add_numbering(doc, bullet=False)
    for s in ['规则层结果必须先过审阅清单：mask --review 的关键信息校验要 ✅，剩余低优先级项逐条确认后再使用。',
              '映射表是敏感文件：含全部原始值，务必加密保存（--encrypt-mapping），切勿上传网络。',
              '云端 LLM 方案数据出境：人名、地址、案情会发送给 API 服务商；涉密材料请用本地 '
              'Ollama 或默认的 Agent 语义层（无需外部 API）。']:
        list_item(doc, warn_id, s)

    # ---------- 已知不足 ----------
    add_heading(doc, '已知不足与边界（不回避）', 1)
    limit_id = add_numbering(doc, bullet=True)
    for s in ['语义层需要"配置了本技能的 AI"来执行：无 AI 环境批量处理时才需本地 Ollama 或云端 API 作为替代；'
              '规则层不承诺覆盖法院名称、案情敏感细节、裸公司简称、OCR 变体写法。',
              '裸人名是启发式：罕见姓氏、单次出现的数字三字名、网络昵称存在漏判或过度脱敏。',
              'jieba 是必装依赖：缺失时裸人名发现自动关闭。',
              '格式保真有限：docx 保留段落结构；PDF 输出为纯文本。',
              'format-converter 只做文本层提取：docx→txt/md 依赖 mammoth（不解析复杂宏/嵌入对象），'
              'pdf→txt 仅限带文本层的电子版文书；扫描件 PDF 需先走 OCR 通道（macOS Vision 内置 / '
              'Windows/Linux 外部 OCR），图片型材料转 PDF 后脱敏效果取决于 OCR 识别率。',
              '仅面向中文法律文书，英文及其他语种实体规则未覆盖。',
              'full 命令依赖模型执行力：不同模型效果有差异，须用 evaluate.py --llm-api 实测。']:
        list_item(doc, limit_id, s)

    # ---------- 对比 ----------
    add_heading(doc, '与同类开源工具的差异', 1)
    make_table(doc,
               ['能力', '本工具', '通用脱敏库', '海外法律脱敏方案'],
               [
                   ['中国法律实体（案号/信用代码/执业证号/车牌）', '✅ 专精', '部分', '❌'],
                   ['语义占位符（[当事人甲（原告）]）', '✅', '❌ 星号打码', '部分'],
                   ['加密映射 + 一键无损还原', '✅', '❌', '部分'],
                   ['量化红队评测（召回率/误报）', '✅', '部分', '部分'],
                   ['全文实体一致（同一人同一占位符）', '✅', '❌', '部分'],
                   ['本地优先 / 数据不出机', '✅', '✅', '✅'],
               ],
               [2.15, 1.45, 1.45, 1.45])

    # ---------- 版本里程碑 ----------
    add_heading(doc, '版本里程碑', 1)
    make_table(doc,
               ['版本', '内容'],
               [
                   ['v1.0', '规则引擎 + LLM 混合架构，18 类实体'],
                   ['v2.1', 'EntityResolver 实体归一化、AES-256-GCM 加密映射表、内存安全模式'],
                   ['v2.1.1', '规则层正确性大修（执业证号/信用代码/身份证/日期/微信边界）'],
                   ['v2.2', 'GB 11643/GB 32100 校验码、一键还原 restore、红队评测 evaluate、本地 NER'],
                   ['v2.3', 'full 完整脱敏流水线（规则层+LLM 合并映射）、失败安全机制'],
                   ['v2.3.1', '云端 API 适配（通义/DeepSeek/智谱），无需本地部署模型'],
                   ['v2.4', '裸人名 + 无层级地址进规则层（jieba 分词过滤 + 全文实体一致）'],
                   ['v2.5', '实战修复：还原精确配对 / 角色词后名词门控 / OCR 空格兼容 / 金额增强；'
                            '59 项回归测试 + 60 个红队用例'],
                   ['v2.6', '两阶段工作流：规则层关键信息必清 → --review 审阅清单 → '
                            'Agent 语义层直接执行（无需外部 API）'],
                   ['v4.1', '内置 format-converter 格式转换核心：docx / pdf / html / csv / 图片 '
                            '→ txt / md / json，纯 JS 无外部引擎（mammoth / pdfjs-dist / '
                            'turndown / marked / csv-parse / sharp），扩充脱敏输入链路；'
                            '来源 FlyingMouse Format（MIT）'],
               ],
               [1.15, 5.35])

    # ---------- 授权 ----------
    add_heading(doc, '授权', 1)
    para(doc, 'MIT License，可自由使用与二次开发。')
    fp = doc.add_paragraph()
    fp.paragraph_format.space_before = Pt(8)
    fp.paragraph_format.space_after = Pt(0)
    r = fp.add_run('本工具是“分层脱敏系统”，不是一键魔法。请根据材料涉密程度选择脱敏深度，'
                   '关键文书建议脱敏后人工抽检。')
    set_run_font(r, size=10, italic=True, color=GRAY)

    # 输出到脚本所在目录（docs/），从仓库根目录执行也能正确落位
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       '简介.docx')
    doc.save(out)
    print('saved', out)


if __name__ == '__main__':
    main()
